#!/usr/bin/env python3
"""
build_switch_relnote_templated.py

Create a templated version of ND-DEL-OBB-035-REL-001-01-Switchv7_8 by cloning the
Nomad Angebot docx (to inherit its branded cover graphic, header/footer logo, fonts,
heading styles and page setup) and pouring in the Switch release-note content.

Method:
  1. Clone the Angebot .docx (preserves cover images, headers/footers, styles, media).
  2. Keep the cover graphics (first body paragraphs that carry the logo/cover image) and
     the final <sectPr>; delete everything in between (the Angebot's offer body).
  3. Rebuild the Switch content with the template's own styles.
  4. Keep the Switch document's identity (title, ref, author Matzner) — NOT the offer text.

Data source: design freeze/ND-DEL-OBB-035-REL-001-01-Switchv7_8.docx (read once; values inlined).
Output:      design freeze/ND-DEL-OBB-035-REL-001-01-Switchv7_8_templated.docx
"""
import copy
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = ".tmp/angebot_template.docx"
OUT = "design freeze/ND-DEL-OBB-035-REL-001-01-Switchv7_8_templated_v2.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)

d = docx.Document(TEMPLATE)
body = d.element.body

# ---------------------------------------------------------------------------
# Step 1: locate the final sectPr (page setup) and the cover graphics to keep.
# Cover = leading paragraphs that contain a <w:drawing> (logo + cover image),
# plus any blank paragraphs before the first real content. From probing:
#   [0],[1] blank, [2],[3] drawings (logo/cover). We keep indices 0..3.
# ---------------------------------------------------------------------------
kids = list(body.iterchildren())
final_sectPr = body.find(qn('w:sectPr'))  # the body-level sectPr (page setup)

# how many leading children to keep: keep through the last drawing-bearing para
# in the first block, then a page break.
keep_until = -1
for i, ch in enumerate(kids):
    if ch.tag.endswith('}p') and ch.findall('.//' + qn('w:drawing')):
        keep_until = i
# keep one trailing blank paragraph after the cover image for spacing
keep_count = keep_until + 1

# Remove every body child from keep_count onward EXCEPT the final sectPr.
for ch in kids[keep_count:]:
    if ch is final_sectPr:
        continue
    body.remove(ch)

# ---------------------------------------------------------------------------
# helpers that insert new elements BEFORE the final sectPr (so it stays last)
# ---------------------------------------------------------------------------
def _style_or_none(name):
    try:
        _ = d.styles[name]; return name
    except KeyError:
        return None

H1 = _style_or_none('Heading 1') or 'Heading 1'
H2 = _style_or_none('Heading 2') or 'Heading 2'
NORMAL = 'Normal'

def add_par(text="", style=None, bold=False, size=None, align=None, color=None, space_before=None, space_after=None):
    p = d.add_paragraph(style=style)
    # move it before the final sectPr (add_paragraph appends at end, after sectPr's parent? -> it's appended to body, before sectPr automatically since sectPr is last child)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size: run.font.size = Pt(size)
        if color is not None: run.font.color.rgb = color
    pf = p.paragraph_format
    if space_before is not None: pf.space_before = Pt(space_before)
    if space_after is not None: pf.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    return p

def add_page_break():
    p = d.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    r._r.append(br)
    return p

def set_cell(cell, text, bold=False, italic=False, size=9, fill=None, color=None, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None: para.alignment = align
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = color
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
        tcPr.append(shd)

def merge_row_fullwidth(table, row_idx):
    """Merge all cells in a row into one (full width)."""
    row = table.rows[row_idx]
    a = row.cells[0]
    for c in row.cells[1:]:
        a = a.merge(c)
    return a

# python-docx appends new paragraphs/tables to the end of the body element.
# Because the body's LAST child is <w:sectPr>, appended block elements land
# BEFORE it automatically (python-docx inserts before a trailing sectPr).

# ---------------------------------------------------------------------------
# Step 2: TITLE BLOCK (document metadata) — reproduce the Switch doc cover data
# ---------------------------------------------------------------------------
title = add_par("DEL-OBB-035 – OEBB DoSto Neu Switch Version 7 und 8",
                style=NORMAL, bold=True, size=18, color=NAVY, space_before=6, space_after=10)

meta = [
    ("Zweck", "Switch Konfiguration Versionskontrolle"),
    ("Herausgegeben an", "OEBB Personenverkehr"),
    ("BMS-Dokumentenreferenz", "BMS-ENGI-FOR-006-02"),
    ("Nomad-Referenz", "ND-DEL-OBB-035-REL-001-01"),
    ("Versionsnummer", "01"),
    ("Letzte Aktualisierung", "22. April 2026"),
    ("Autor", "Johannes Matzner"),
    ("Freigabe durch", "Derek Abdinor"),
    ("Anzahl der Seiten", "5"),
    ("Zugehörige Dokumente", "keine"),
]
tbl_meta = d.add_table(rows=len(meta), cols=2)
tbl_meta.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    set_cell(tbl_meta.rows[i].cells[0], k, bold=True, size=10, fill="D6E0E5")
    set_cell(tbl_meta.rows[i].cells[1], v, size=10)
tbl_meta.columns[0].width = Pt(150)

add_par("", style=NORMAL)
# Disclaimer
add_par("Haftungsausschluss", style=H2)
add_par(
    "Dieses Dokument sowie alle zugehörigen Dokumente enthalten streng vertrauliche "
    "Informationen und sind ausschließlich für die vorgesehenen Empfänger bestimmt. "
    "Jegliche unbefugte Offenlegung, Weitergabe, Verbreitung oder Vervielfältigung dieses "
    "Dokuments oder zugehöriger Dokumente – in welcher Form auch immer – ist strengstens untersagt.",
    style=NORMAL)

# Document version control (the doc's own revision table)
add_par("Versionskontrolle", style=H2)
docctl = [
    ("Version", "Revision date", "Summary of edits", "Created By", "Approved By"),
    ("1", "22.04.2026", "Erste Entwurf", "J. Matzner", "D. Abdinor"),
]
t0 = d.add_table(rows=len(docctl), cols=5); t0.style = 'Table Grid'
for j, h in enumerate(docctl[0]):
    set_cell(t0.rows[0].cells[j], h, bold=True, size=9, fill="2E6171", color=RGBColor(0xFF,0xFF,0xFF))
for j, v in enumerate(docctl[1]):
    set_cell(t0.rows[1].cells[j], v, size=9)

add_page_break()

# ---------------------------------------------------------------------------
# Step 3: TABLE OF CONTENTS (auto-updating field)
# ---------------------------------------------------------------------------
add_par("Inhalt", style=_style_or_none('TOC Heading') or H1)
toc_p = d.add_paragraph()
run = toc_p.add_run()
fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = r'TOC \o "1-3" \h \z \u'
fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
placeholder = OxmlElement('w:r'); pt = OxmlElement('w:t'); pt.text = "Rechtsklick → Felder aktualisieren, um das Inhaltsverzeichnis zu erzeugen."; placeholder.append(pt)
fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep)
toc_p._p.append(placeholder);
run2 = toc_p.add_run(); run2._r.append(fldEnd)

add_page_break()

# ---------------------------------------------------------------------------
# Step 4: §1 EINLEITUNG
# ---------------------------------------------------------------------------
add_par("1. Einleitung", style=H1)
add_par("Zweck", style=H2)
add_par(
    "Dosto Neu-Flotten enthalten wichtige Netzwerkeinstellungen in den Switch-Konfigurationen. "
    "Da Designänderungen Auswirkungen auf die Flottenkonfiguration haben, dokumentiert dieses "
    "Dokument die genehmigten Versionen und die wichtigsten Änderungen, die vorgenommen wurden, "
    "um die Betriebs- und Geschäftskontinuität zu gewährleisten.", style=NORMAL)

add_par("Reifegrad", style=H2)
add_par("Alle in dieser Version des Dokuments beschriebenen Softwarestände werden geliefert für:", style=NORMAL)
maturity = [
    ("Tests auf dem Zug", "Freigegeben"),
    ("Vor-kommerzieller Betrieb (Bereitstellung auf ausgelieferten Zügen)", "Freigegeben"),
    ("Kundentests auf dem Zug", "Freigegeben"),
    ("Kommerzieller Betrieb (Rollout auf die Flotte)", "Freigegeben"),
]
tm = d.add_table(rows=len(maturity)+1, cols=2); tm.style = 'Table Grid'
set_cell(tm.rows[0].cells[0], "Phase", bold=True, size=9, fill="2E6171", color=RGBColor(0xFF,0xFF,0xFF))
set_cell(tm.rows[0].cells[1], "Status", bold=True, size=9, fill="2E6171", color=RGBColor(0xFF,0xFF,0xFF))
for i,(ph,stt) in enumerate(maturity, start=1):
    set_cell(tm.rows[i].cells[0], ph, size=9)
    set_cell(tm.rows[i].cells[1], stt, size=9, color=RGBColor(0x54,0x82,0x35))

# ---------------------------------------------------------------------------
# Step 5: §2 VERSIONSKONTROLLE DER SWITCHES (per-fleet V5..V8)
# ---------------------------------------------------------------------------
add_par("2. Versionskontrolle der Switches", style=H1)
add_par(
    "In diesem Kapitel werden die Software-Release-Versionen der OEBB DOSTO NEU IOB-Switch-"
    "Konfigurationen sowie die daran vorgenommenen Änderungen dokumentiert. Die Switch-"
    "Konfigurationen werden je Flottentyp aus einer gemeinsamen Vorlage generiert; die "
    "inhaltlichen Änderungen je Version sind flottenübergreifend identisch, mit geringfügigen "
    "mechanischen Abweichungen je Wagenanzahl (siehe je Flotte vermerkt).", style=NORMAL)

# ---- Bullets per version (shared change content) ----
V5_B = [
    "Für die Sprechstellen in Wagen C wurden die DHCP-Parameter korrigiert.",
    "Die AFZ-Ports DHCP-Parameter wurden vervollständigt.",
    "Die Bistro-Ports wurden konfiguriert.",
    "Deaktivieren des Nomad-Service-Ports (Security).",
]
V6_B = [
    "Es wurde DNS-Information (8.8.8.8) an Port-Based DHCP hinzugefügt.",
]
V7_B = [
    "DNS und NTP wurden auf RCU 192.168.101.10 geändert.",
    "Frontkupplung auf VLANs 3, 5, 9 gesetzt.",
]
V8_B = [
    "RSTP (Rapid Spanning Tree Protocol) wurde implementiert. Die RSTP-Kosten für die "
    "Multitraction-Ports wurden manuell angepasst (bei nv4 dynamisch über die train_id), um die "
    "Verfügbarkeit des gesamten Netzwerks sicherzustellen.",
    "RSTP-Prioritäten wurden ergänzt; die Verbindung B1–B3 ist als bevorzugter Punkt definiert, "
    "um den Ring in eine stabile U-Topologie zu überführen.",
    "Die Mehrfachtraktions-Ports an den Switches A3 und B3 werden gegenüber A1 und B1 priorisiert.",
    "VLAN 15 wurde den Firewall- und Multitraction-Ports hinzugefügt. Die Mehrfachtraktion führt "
    "ausschließlich VLAN 5 und VLAN 15.",
    "Der Port e2-5 auf den Switches C3 und F3 wurde für FV6 konfiguriert (FV6 ist die einzige "
    "Einheit mit Bistro).",
    "Korrektur überlappender Access-Point-Frequenzen (nv6: ab Wagen 3; nv4: gespiegelte AP-"
    "Konfiguration ab Wagen 2).",
    "Spanning-tree Edge-Modus manuell gesetzt; fis/service/sprechstelle auf /25-Netze geändert.",
]

# (version, date, by, approved, bullets)
ALL_RELEASES = {
    "V5": ("11–13.11.2025", "Davud Zejnelovic", "Davud Zejnelovic", V5_B),
    "V6": ("10–15.12.2025", "Davud Zejnelovic", "Davud Zejnelovic", V6_B),
    "V7": ("16.02.2026",   "Davud Zejnelovic", "Davud Zejnelovic", V7_B),
    "V8": ("31.03–20.04.2026", "Davud Zejnelovic", "Davud Zejnelovic", V8_B),
}

# ---- Fleet applicability matrix (verified) ----
add_par("2.1 Versions- und Flottenmatrix", style=H2)
add_par(
    "Die folgende Matrix zeigt, welche Konfigurationsversion für welchen Flottentyp freigegeben "
    "wurde. Stand der Verifikation: Konfigurations-Repositories (nv4, nv6) sowie gerenderte "
    "Konfigurations-Snapshots.", style=NORMAL)
mtx_hdr = ["Version", "4734 (nv4, 4-Teiler)", "4736 (nv6, 6-Teiler)", "4706 (fv6, FV 6-Teiler)", "4705 (fv5, CAT 5-Teiler)"]
# ✓ = verified present; – = not applicable/not yet; (?) = to confirm
mtx_rows = [
    ("V5", "✓", "✓", "✓", "–  (erst ab V6)"),
    ("V6", "✓", "✓", "✓", "✓"),
    ("V7", "✓", "✓", "✓", "✓"),
    ("V8", "✓", "✓", "✓", "✓"),
]
mt = d.add_table(rows=len(mtx_rows)+1, cols=5); mt.style = 'Table Grid'
for j,h in enumerate(mtx_hdr):
    set_cell(mt.rows[0].cells[j], h, bold=True, size=9, fill="2E6171", color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
for i,row in enumerate(mtx_rows, start=1):
    for j,v in enumerate(row):
        set_cell(mt.rows[i].cells[j], v, bold=(j==0), size=9, align=(WD_ALIGN_PARAGRAPH.CENTER if j>0 else None))
add_par(
    "Legende:  ✓ = freigegeben und verifiziert (Konfigurations-Repository)   ·   "
    "– = nicht zutreffend / noch nicht eingeführt.   "
    "Verifiziert am 01.06.2026 gegen die vier OBN-Konfigurations-Repositories (nv4, nv6, fv5, fv6).",
    style=NORMAL, size=8)
add_par(
    "Hinweis: Der CAT-5-Teiler (4705 / fv5) wurde erst ab Version V6 in die Flotte eingeführt; "
    "für diesen Flottentyp existiert keine V5. Alle vier Flottentypen befinden sich aktuell auf V8.",
    style=NORMAL, bold=True)

# ---- helper to render a version block (table + bullets) within a fleet ----
def version_block(ver, fleet_note=None):
    datum, erst, abg, bullets = ALL_RELEASES[ver]
    add_par(f"Switch-Konfiguration {ver}", style=_style_or_none('Heading 3') or H2)
    t = d.add_table(rows=2, cols=4); t.style = 'Table Grid'
    for j, h in enumerate(["Version", "Erstellt am", "Erstellt von", "Abgenommen von"]):
        set_cell(t.rows[0].cells[j], h, bold=True, size=9, fill="2E6171", color=RGBColor(0xFF,0xFF,0xFF))
    for j, v in enumerate([ver, datum, erst, abg]):
        set_cell(t.rows[1].cells[j], v, size=9)
    add_par("Neue Funktionen und behobene Probleme", style=NORMAL, bold=True, space_before=4, space_after=2)
    bl = list(bullets) + ([fleet_note] if fleet_note else [])
    for b in bl:
        p = add_par(b, style=_style_or_none('List Bullet') or NORMAL)

# ---- per-fleet sections ----
FLEETS = [
    ("2.2 Flotte 4734 — Nahverkehr 4-Teiler (nv4)",
     "Konfigurationsvorlage: nd-obn-template-dostoneu-nv4. Verifizierte Versionen: V5–V8.",
     ["V5","V6","V7","V8"],
     {"V8": "nv4-spezifisch: STP-Kosten werden dynamisch über die train_id gesetzt; "
            "gespiegelte AP-Konfiguration (x1–x3, x2–x4) ab Wagen 2."}),
    ("2.3 Flotte 4736 — Nahverkehr 6-Teiler (nv6)",
     "Konfigurationsvorlage: nd-obn-template-dostoneu-nv6. Verifizierte Versionen: V5–V8.",
     ["V5","V6","V7","V8"],
     {"V8": "nv6-spezifisch: Korrektur überlappender AP-Frequenzen ab Wagen 3."}),
    ("2.4 Flotte 4706 — Fernverkehr 6-Teiler (fv6)",
     "Konfigurationsvorlage: nd-obn-template-dostoneu-fv6. Verifizierte Versionen: V5–V8. "
     "Einziger Flottentyp mit Bistro (Port e2-5 auf C3/F3).",
     ["V5","V6","V7","V8"],
     {"V8": "fv6-spezifisch: Port e2-5 auf C3/F3 für das Bistro konfiguriert (nur FV6)."}),
    ("2.5 Flotte 4705 — CAT/Fernverkehr 5-Teiler (fv5)",
     "Konfigurationsvorlage: nd-obn-template-dostoneu-fv5. Eingeführt ab V6; verifizierte "
     "Versionen: V6–V8. Keine V5.",
     ["V6","V7","V8"],
     {}),
]
for title, note_txt, vers, fleet_notes in FLEETS:
    add_par(title, style=H2)
    add_par(note_txt, style=NORMAL)
    for v in vers:
        version_block(v, fleet_notes.get(v))

# ---------------------------------------------------------------------------
# Step 6: §4 EIGENSCHAFTEN
# ---------------------------------------------------------------------------
add_par("3. Eigenschaften", style=H1)
add_par("Neue Funktionalitäten", style=H2)
add_par("Wie oben für jede Version aufgeführt.", style=NORMAL)
add_par("Auslassungen und Einschränkungen", style=H2)
add_par("k. A.", style=NORMAL)
add_par("Medieninhalte", style=H2)
add_par("k. A.", style=NORMAL)
add_par("Tests und Validierung", style=H2)
add_par(
    "Die initialen Funktionstests wurden auf dem Testbench durchgeführt. Anschließend erfolgten "
    "die Abnahme durch Stadler sowie Tests auf dem Zug. Nach erfolgreicher Abnahme wurde die "
    "Freigabe für den Rollout in die Flotte erteilt.", style=NORMAL)

d.save(OUT)
print("WROTE", OUT)
