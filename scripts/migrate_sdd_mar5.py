#!/usr/bin/env python3
"""
migrate_sdd_mar5.py — Purge MAR3 / 10.178 / 172.16.195 from the SDD, replace with
MAR5 / 10.179 / TUN-IPs. Input is the comment-stripped clean copy; output is the
customer-ready copy.

Kept deliberately: Engineering-'Actions' page text "MAR3" + MQTT topic "mar3/action"
(real software identifiers, per instruction).

Stage A: prose string-edits on word/document.xml (single zip pass -> temp).
Stage B: rebuild Table 7 (rows) + Table 8 (5-server MAR5 list) via python-docx.

In : design freeze/ND-DEL-OBB-035-SDD-002-01_v2.2_clean.docx
Out: design freeze/ND-DEL-OBB-035-SDD-002-01_v2.2_customer.docx
"""
import os, zipfile
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "design freeze/ND-DEL-OBB-035-SDD-002-01_v2.2_clean.docx"
TMP = "design freeze/.sdd_mar5_tmp.docx"
OUT = "design freeze/ND-DEL-OBB-035-SDD-002-01_v2.2_customer.docx"

MAR5 = [
    ("vmmar5be31.obb.21net.com", "77.237.62.210"),
    ("vmmar5be32.obb.21net.com", "77.237.62.212"),
    ("vmmar5be33.obb.21net.com", "77.237.62.214"),
    ("vmmar5be34.obb.21net.com", "77.237.62.216"),
    ("vmmar5be35.obb.21net.com", "77.237.62.218"),
]

# ---------------- Stage A: prose ----------------
def edit_prose(xml: str) -> str:
    reps = [
        ("Das Backend ist in zwei IP-Adressbereiche unterteilt, von denen jeder bis zu ",
         "Das Backend nutzt für die Dosto-Neu-Flotte (Nomad Connect Project ID 51) den Adressbereich "),
        ("<w:t>127 CCUs</w:t>", "<w:t>10.179.0.0/16</w:t>"),
        ("<w:t>10.178.0.0/17</w:t>", "<w:t>Jeder Zug erhält das Subnetz 10.179.&lt;train-id&gt;.0/24.</w:t>"),
        ("<w:t>10.179.0.0/17</w:t>", "<w:t xml:space=\"preserve\"> </w:t>"),
        ("<w:t>254 CCUs</w:t>", "<w:t>bis zu 255 Fahrzeuge (8-Bit-Fahrzeug-ID)</w:t>"),
        ("<w:t>sechs Home Agents</w:t>", "<w:t>fünf MAR5 Home Agents</w:t>"),
        (" eingesetzt, von denen jeder 14 CCUs hostet.",
         " eingesetzt (ein weiterer, sechster Home Agent ist in Bereitstellung)."),
    ]
    for a, b in reps:
        xml = xml.replace(a, b)
    return xml

with zipfile.ZipFile(SRC) as zin:
    with zipfile.ZipFile(TMP, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = edit_prose(data.decode("utf-8")).encode("utf-8")
            zi = zipfile.ZipInfo(item.filename, date_time=item.date_time)
            zi.compress_type = item.compress_type
            zi.external_attr = item.external_attr
            zout.writestr(zi, data)

# ---------------- Stage B: tables ----------------
d = docx.Document(TMP)

def set_cell(cell, text, bold=False, size=9, fill=None, white=False):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

# Table 7
t7 = d.tables[7]
set_cell(t7.rows[1].cells[1], "10.179.0.0/16")
set_cell(t7.rows[3].cells[0], "MAR5 Backend-Server")
set_cell(t7.rows[3].cells[1], "77.237.62.210 – 77.237.62.218")
set_cell(t7.rows[3].cells[2], "MAR5 Home Agents (vmmar5be31 – vmmar5be35; ein weiterer in Bereitstellung)")
set_cell(t7.rows[4].cells[1], "10.179.13.1")

# Table 8 — rebuild as MAR5 server list (drop per-train ranges)
t8 = d.tables[8]
tbl = t8._tbl
ncols = len(t8.columns)
for row in list(t8.rows):
    tbl.remove(row._tr)

def add_row(vals, header=False):
    row = t8.add_row()
    padded = vals + [""] * (ncols - len(vals))
    for j, v in enumerate(padded):
        set_cell(row.cells[j], v, bold=header, fill=("2E6171" if header else None), white=header)
    return row

add_row(["MAR5 Home Agent", "TUN-IP"], header=True)
for name, ip in MAR5:
    add_row([name, ip])

# explanatory note after Table 8
note = OxmlElement('w:p')
note.append(OxmlElement('w:pPr'))
r = OxmlElement('w:r')
rpr = OxmlElement('w:rPr'); rpr.append(OxmlElement('w:i'))
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz)
r.append(rpr)
t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
t.text = ("Die Zuordnung der Fahrzeuge zu den Home Agents erfolgt gemäß dem Nomad Connect "
          "Project-ID-51-Schema (Fahrzeug-Subnetz 10.179.<train-id>.0/24). Aktuell sind fünf "
          "MAR5 Home Agents bereitgestellt; ein sechster ist in Bereitstellung.")
r.append(t); note.append(r)
tbl.addnext(note)

d.save(OUT)
os.remove(TMP)
print("WROTE", OUT)
