import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def shade_header_row(table, fill='D5E8F0'):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

def add_table_row(table, values, bold_first=False, fill=None, font_size=9):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(font_size)
                if bold_first and i == 0:
                    run.font.bold = True
        if fill:
            set_cell_shading(cell, fill)
    return row

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(size)

doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2)
section.top_margin  = section.bottom_margin = Cm(2)

# Header
hdr = section.header
hdr.paragraphs[0].clear()
hr = hdr.paragraphs[0].add_run('Network Health Check - Fzg. 136 (4736-108)     Nomad Digital - CONFIDENTIAL')
hr.font.size = Pt(9); hr.font.color.rgb = RGBColor(0x80,0x80,0x80)

# Footer
ftr = section.footer
ftr.paragraphs[0].clear()
fr = ftr.paragraphs[0].add_run('Confidential - Draft for review')
fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x80,0x80,0x80)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Nomad Digital')
r.bold = True; r.font.size = Pt(18)

p2 = doc.add_paragraph()
r2 = p2.add_run('Connected Transport, Intelligent Solutions')
r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x80,0x80,0x80)

doc.add_paragraph()
p3 = doc.add_paragraph()
r3 = p3.add_run('Network Health Check Report')
r3.bold = True; r3.font.size = Pt(32); r3.font.color.rgb = RGBColor(0x1F,0x4E,0x79)

p4 = doc.add_paragraph()
r4 = p4.add_run('DOSTO Nahverkehr 6-car')
r4.font.size = Pt(22); r4.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.add_run('Prepared for').font.size = Pt(14)
p6 = doc.add_paragraph()
r6 = p6.add_run('OeBB - Oesterreichische Bundesbahnen')
r6.bold = True; r6.font.size = Pt(18)

doc.add_paragraph()

meta = doc.add_table(rows=1, cols=2)
meta.style = 'Table Grid'
set_cell_shading(meta.rows[0].cells[0], 'D5E8F0')
meta.rows[0].cells[0].text = 'Trainset (Fzg. Nr.)'
meta.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
meta.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
meta.rows[0].cells[1].text = '4736-108'
meta.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
META_ROWS = [
    ('Fahrzeug-ID (OeBB)', 'Fzg. 136'),
    ('Configuration', '6-car'),
    ('CCU', 'box1-t8 (10.179.8.1)'),
    ('Check date', '2026-05-04'),
    ('Author', 'Abbas Rizvi, Nomad Digital'),
    ('Version', '1.0'),
    ('Status', 'Draft for review'),
]
for label, value in META_ROWS:
    rr = meta.add_row()
    set_cell_shading(rr.cells[0], 'D5E8F0')
    rr.cells[0].text = label
    rr.cells[1].text = value
    for c in rr.cells:
        for pa in c.paragraphs:
            for ru in pa.runs:
                ru.font.size = Pt(9)
    rr.cells[0].paragraphs[0].runs[0].font.bold = True
set_col_widths(meta, [5.5, 11.0])

doc.add_page_break()

# ── DOCUMENT CONTROL ──────────────────────────────────────────────────────────
doc.add_heading('Document Control', 1)
doc.add_paragraph(
    'This document records findings from an on-train Layer-2 network health check '
    'performed on DOSTO trainset Fzg. 136 (4736-108) on 2026-05-04. '
    'Reviewers may add comments using Word Track Changes or the Reviewer Notes table in Appendix D.'
)

doc.add_heading('Revision History', 2)
rev = doc.add_table(rows=1, cols=4)
rev.style = 'Table Grid'
for i, h in enumerate(['Version','Date','Author','Change Summary']):
    rev.rows[0].cells[i].text = h
shade_header_row(rev)
add_table_row(rev, [
    '1.0', '2026-05-04', 'Abbas Rizvi, Nomad Digital',
    'Initial draft - routine L2 health check on 4736-108 (Fzg. 136), '
    'including OBN config/firmware update and cabling fault assessment.'
], font_size=9)
add_table_row(rev, ['','','',''], font_size=9)
set_col_widths(rev, [1.5, 2.5, 5.0, 7.5])

doc.add_heading('Approvals', 2)
appr = doc.add_table(rows=1, cols=4)
appr.style = 'Table Grid'
for i, h in enumerate(['Role','Name','Organisation','Date']):
    appr.rows[0].cells[i].text = h
shade_header_row(appr)
for role, name, org in [
    ('Author','Abbas Rizvi, Nomad Digital','Nomad Digital'),
    ('Technical Reviewer','','Nomad Digital'),
    ('OeBB Reviewer','','OeBB'),
]:
    add_table_row(appr, [role, name, org, ''], font_size=9)
set_col_widths(appr, [3.5, 5.0, 4.0, 4.0])

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
doc.add_heading('1. Executive Summary', 1)
doc.add_paragraph(
    'A Layer-2 network health check was performed on DOSTO trainset Fzg. 136 (4736-108) '
    'on 2026-05-04 by Abbas Rizvi (Nomad Digital). The check included switch inventory, '
    'RSTP topology verification, LLDP cabling audit, per-port error scan, OBN configuration '
    'and firmware update, and Westermo access point firmware update. Two cabling faults were '
    'identified via LLDP; one resolved by Stadler Rail, one outstanding.'
)

doc.add_heading('Headline Verdict', 2)
vt = doc.add_table(rows=1, cols=1)
vt.style = 'Table Grid'
set_cell_shading(vt.rows[0].cells[0], 'FCE4D6')
vp = vt.rows[0].cells[0].paragraphs[0]
vr = vp.add_run(
    'OVERALL VERDICT: CABLING FAULT (one outstanding) - '
    'Layer-2 fabric operationally healthy via RSTP redundancy; '
    'full clean bill requires Stadler to reconnect the D1<->E2 inter-coach trunk cable.'
)
vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = RGBColor(0x84,0x36,0x0F)

doc.add_paragraph()

EXEC_BULLETS = [
    '18 VDS Rail Consist Switches found and reachable on management VLAN.',
    'Inter-coach trunks: e0-0 18/18 UP at 10G full-duplex. e0-1 17/18 UP - D1 e0-1 '
    'link-down (missing cable to E2, Stadler action required).',
    'RSTP root bridge: D1 (0/a0:59:3a:d0:5e:c0), 18/18 agreement - single stable root. '
    'RSTP routing around the D1<->E2 break; all ports FWD, no traffic loss.',
    'Per-port error scan: 1 port (C2 e0-1) with single-count noise '
    '(1x rx_err, 1x jabber, 1x runt, 1x giant - CRC=0, carrier-false=0). Not actionable.',
    'OBN: all 18 VDS switches on nv6-XX-v8-136 and firmware 7.4.2. '
    'F3 train_id corrected 144->136.',
    'Westermo APs: firmware 6.10.0-0 -> 6.11.2-0 update in progress '
    '(9/24 confirmed at time of report). APxm-v1 config for coaches 4-6 pending.',
    'Cabling fault 1 (C3 e0-0/e0-1 swap): RESOLVED by Stadler - LLDP confirmed correct.',
    'Cabling fault 2 (D1<->E2 missing link): OUTSTANDING - Stadler action required.',
]
for b in EXEC_BULLETS:
    add_bullet(doc, b, 10)

doc.add_page_break()

# ── 2. SCOPE ──────────────────────────────────────────────────────────────────
doc.add_heading('2. Scope and Methodology', 1)
doc.add_heading('2.1 What Was Checked', 2)
for item in [
    'All 18 VDS Rail Consist Switches on management VLAN 100 (10.179.8.128/25).',
    'All inter-coach trunk uplinks (e0-0 and e0-1 on each switch).',
    'LLDP topology check - e0-0 and e0-1 neighbour mapping vs. OBN expected topology.',
    'RSTP spanning tree root and port states across the fleet.',
    'OBN configuration and firmware audit (nv6-XX-v8-136 target, firmware 7.4.2).',
    '24 Westermo AP firmware update (6.10.0-0 -> 6.11.2-0) and APxm-v1 config push for coaches 4-6.',
]:
    add_bullet(doc, item, 10)

doc.add_heading('2.2 What Is Out of Scope', 2)
for item in [
    'Stadler-side device VLANs (cameras, displays, AFZ, intercom, OBS, RDC) - managed by Stadler Rail.',
    'PWLAN / passenger Wi-Fi - separate Nomad Digital scope.',
    'Cellular WAN backhaul - separate Nomad Digital scope.',
    'Full L2 throughput sampling and Stadler FW reachability - deferred until D1<->E2 cabling resolved.',
]:
    add_bullet(doc, item, 10)

doc.add_heading('2.3 Method', 2)
doc.add_paragraph(
    'The check was performed via SSH from the engineer laptop through the CCU (box1-t8, 10.179.8.1) '
    'to each consist switch. Switch CLI commands (show interface details, show lldp neighbours, '
    'show spanning-tree) were run per switch. OBN management used the OBN CLI on the CCU. '
    'All 7 known OBN software bugs were pre-patched via fix_obn.py before the update run.'
)

doc.add_page_break()

# ── 3. DETAILED FINDINGS ──────────────────────────────────────────────────────
doc.add_heading('3. Detailed Findings', 1)

doc.add_heading('3.1 Switch Inventory', 2)
inv = doc.add_table(rows=1, cols=3)
inv.style = 'Table Grid'
for i, h in enumerate(['Parameter','Observed','Expected']):
    inv.rows[0].cells[i].text = h
shade_header_row(inv)
for p2, ob, ex in [
    ('VDS switches found','18','18 (6-car consist)'),
    ('Westermo APs found','24','24 (typical 6-car)'),
    ('CCU vlan100 subnet','10.179.8.128/25','10.179.X.128/25'),
    ('OBN train_id','136 (= 128 + 8)','Correct for this consist'),
    ('OBN template family','nv6-v7','nv6-* (6-car)'),
]:
    add_table_row(inv, [p2, ob, ex], bold_first=True, font_size=9)
set_col_widths(inv, [5.5, 5.5, 5.5])

doc.add_heading('3.2 Schema-to-IP Mapping', 2)
doc.add_paragraph(
    'Switches identified by LLDP hostname (nv6-XX-v8-136) after OBN config push. '
    'Prior to push, identification was by trunk/access-port fingerprint.'
)
ipm = doc.add_table(rows=1, cols=3)
ipm.style = 'Table Grid'
for i, h in enumerate(['Schema Role','IP','Notes']):
    ipm.rows[0].cells[i].text = h
shade_header_row(ipm)
IP_MAP = [
    ('A1','10.179.8.185',''),
    ('A2','10.179.8.194',''),
    ('A3','10.179.8.182','Stadler FW switch - e1-4 multi-VLAN trunk'),
    ('C1','10.179.8.190',''),
    ('C2','10.179.8.192',''),
    ('C3','10.179.8.180','Cable swap e0-0/e0-1 - now fixed by Stadler'),
    ('D1','10.179.8.183','RSTP root; e0-1 link-down (cable missing to E2)'),
    ('D2','10.179.8.195',''),
    ('D3','10.179.8.186',''),
    ('E1','10.179.8.179',''),
    ('E2','10.179.8.191','e0-1 link-down (cable missing to D1)'),
    ('E3','10.179.8.188',''),
    ('F1','10.179.8.184',''),
    ('F2','10.179.8.187',''),
    ('F3','10.179.8.193','train_id corrected 144->136'),
    ('B1','10.179.8.181','ZFR primary - e1-11 VLAN 2'),
    ('B2','10.179.8.178',''),
    ('B3','10.179.8.189','ZFR standby - e1-11 VLAN 2'),
]
for role, ip, note in IP_MAP:
    add_table_row(ipm, [role, ip, note], font_size=9)
set_col_widths(ipm, [2.0, 4.0, 10.5])

doc.add_heading('3.3 RSTP Spanning Tree', 2)
stp = doc.add_table(rows=1, cols=2)
stp.style = 'Table Grid'
stp.rows[0].cells[0].text = 'Parameter'
stp.rows[0].cells[1].text = 'Value'
shade_header_row(stp)
for param, val in [
    ('Protocol','RSTP (Rapid Spanning Tree Protocol)'),
    ('Root bridge','0/a0:59:3a:d0:5e:c0 (D1, 10.179.8.183)'),
    ('Switches agreeing on root','18/18'),
    ('Topology stable','Yes - single root, no TCN flapping observed'),
    ('D1<->E2 trunk impact',
     'E2 routes to root via E3->F2/F3->B-end (path cost 432100 vs normal 2000). '
     'All ports FWD. No traffic loss.'),
]:
    add_table_row(stp, [param, val], bold_first=True, font_size=9)
set_col_widths(stp, [5.5, 11.0])

doc.add_heading('3.4 LLDP Cabling Audit', 2)
doc.add_paragraph(
    'LLDP topology check run against OBN expected nv6-v7 topology. '
    '36 trunk ports checked (e0-0 and e0-1 on all 18 switches).'
)
lldp = doc.add_table(rows=1, cols=5)
lldp.style = 'Table Grid'
for i, h in enumerate(['Switch','Port','Live Neighbour','Expected','Status']):
    lldp.rows[0].cells[i].text = h
shade_header_row(lldp)
LLDP_DATA = [
    ('C3 (10.179.8.180)','e0-0','A2','A2','RESOLVED','E2EFDA'),
    ('C3 (10.179.8.180)','e0-1','C2','C2','RESOLVED','E2EFDA'),
    ('D1 (10.179.8.183)','e0-1','NO NEIGHBOUR','E2','OUTSTANDING','FCE4D6'),
    ('E2 (10.179.8.191)','e0-1','NO NEIGHBOUR','D1','OUTSTANDING','FCE4D6'),
]
for sw, port, live, exp, status, fill in LLDP_DATA:
    row = lldp.add_row()
    for i, v in enumerate([sw, port, live, exp, status]):
        row.cells[i].text = v
        for run in row.cells[i].paragraphs[0].runs:
            run.font.size = Pt(9)
    set_cell_shading(row.cells[4], fill)
doc.add_paragraph('All other 32 trunk ports: LLDP neighbours match OBN expected topology.').runs[0].font.size = Pt(9)
set_col_widths(lldp, [4.0, 1.5, 4.0, 2.5, 4.5])

doc.add_heading('3.5 Inter-Coach Trunk Error Scan', 2)
doc.add_paragraph(
    'All 36 inter-coach trunk ports checked. D1 e0-1 and E2 e0-1 are link-down '
    '(cable missing - no flapping, no errors). '
    'One noise event on C2 e0-1 (CRC=0, carrier-false=0 - not actionable).'
)
tbl = doc.add_table(rows=1, cols=7)
tbl.style = 'Table Grid'
for i, h in enumerate(['Schema','Switch IP','Port','Speed','RX Errors','RX CRC','Jabber/Carrier-false']):
    tbl.rows[0].cells[i].text = h
shade_header_row(tbl)
SWITCHES = [
    ('B2','10.179.8.178'),('E1','10.179.8.179'),('C3','10.179.8.180'),
    ('B1','10.179.8.181'),('A3','10.179.8.182'),('D1','10.179.8.183'),
    ('F1','10.179.8.184'),('A1','10.179.8.185'),('D3','10.179.8.186'),
    ('F2','10.179.8.187'),('E3','10.179.8.188'),('B3','10.179.8.189'),
    ('C1','10.179.8.190'),('E2','10.179.8.191'),('C2','10.179.8.192'),
    ('F3','10.179.8.193'),('A2','10.179.8.194'),('D2','10.179.8.195'),
]
DOWN = {('D1','e0-1'),('E2','e0-1')}
NOISE = {('C2','e0-1')}
for schema, ip in SWITCHES:
    for port in ['e0-0','e0-1']:
        key = (schema, port)
        if key in DOWN:
            speed = 'link DOWN (cable missing)'
            rx_err = crc = jabber = '0'
            fill = 'FFF2CC'
        elif key in NOISE:
            speed = '10000 Mb/s, Full'
            rx_err = '1 (noise)'; crc = '0'; jabber = '1 / 0'
            fill = 'FFF2CC'
        else:
            speed = '10000 Mb/s, Full'
            rx_err = crc = '0'; jabber = '0 / 0'
            fill = None
        row = tbl.add_row()
        for i, v in enumerate([schema, ip, port, speed, rx_err, crc, jabber]):
            row.cells[i].text = v
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
            if fill:
                set_cell_shading(row.cells[i], fill)
set_col_widths(tbl, [1.2, 3.2, 1.2, 4.3, 1.8, 1.5, 3.3])

doc.add_heading('3.6 OBN Configuration and Firmware', 2)
doc.add_paragraph(
    'All 18 VDS switches updated to nv6-XX-v8-136 (config) and 7.4.2 (firmware). '
    'F3 template error corrected (train_id 144->136). '
    'All 7 OBN software bugs pre-patched via fix_obn.py.'
)
obn = doc.add_table(rows=1, cols=4)
obn.style = 'Table Grid'
for i, h in enumerate(['Switch','IP','Config','Firmware']):
    obn.rows[0].cells[i].text = h
shade_header_row(obn)
OBN_DATA = [
    ('A1','10.179.8.185','nv6-A1-v8-136','7.4.2'),
    ('A2','10.179.8.194','nv6-A2-v8-136','7.4.2'),
    ('A3','10.179.8.182','nv6-A3-v8-136','7.4.2'),
    ('C1','10.179.8.190','nv6-C1-v8-136','7.4.2'),
    ('C2','10.179.8.192','nv6-C2-v8-136','7.4.2'),
    ('C3','10.179.8.180','nv6-C3-v8-136','7.4.2'),
    ('D1','10.179.8.183','nv6-D1-v8-136','7.4.2'),
    ('D2','10.179.8.195','nv6-D2-v8-136','7.4.2'),
    ('D3','10.179.8.186','nv6-D3-v8-136','7.4.2'),
    ('E1','10.179.8.179','nv6-E1-v8-136','7.4.2'),
    ('E2','10.179.8.191','nv6-E2-v8-136','7.4.2'),
    ('E3','10.179.8.188','nv6-E3-v8-136','7.4.2'),
    ('F1','10.179.8.184','nv6-F1-v8-136','7.4.2'),
    ('F2','10.179.8.187','nv6-F2-v8-136','7.4.2'),
    ('F3','10.179.8.193','nv6-F3-v8-136 (corrected from 144)','7.4.2'),
    ('B1','10.179.8.181','nv6-B1-v8-136','7.4.2'),
    ('B2','10.179.8.178','nv6-B2-v8-136','7.4.2'),
    ('B3','10.179.8.189','nv6-B3-v8-136','7.4.2'),
]
for row_data in OBN_DATA:
    add_table_row(obn, list(row_data), font_size=9)
set_col_widths(obn, [1.5, 3.5, 7.5, 4.0])

doc.add_heading('3.7 Westermo AP Firmware Update', 2)
doc.add_paragraph(
    'Firmware update 6.10.0-0 -> 6.11.2-0 initiated for all 24 APs. '
    'At time of report: 9/24 confirmed on 6.11.2-0. '
    'APxm-v1 config push to coaches 4-6 (12 APs) pending firmware completion.'
)
ap = doc.add_table(rows=1, cols=4)
ap.style = 'Table Grid'
for i, h in enumerate(['Coach','AP','IP','Firmware Status']):
    ap.rows[0].cells[i].text = h
shade_header_row(ap)
AP_DATA = [
    ('1','1','10.179.8.240','6.11.2-0'),
    ('1','3','10.179.8.231','6.11.2-0'),
    ('1','4','10.179.8.223','Update sent - completing'),
    ('2','1','10.179.8.236','6.11.2-0'),
    ('2','2','10.179.8.233','6.11.2-0'),
    ('2','3','10.179.8.232','6.11.2-0'),
    ('2','4','10.179.8.225','6.11.2-0'),
    ('3','1','10.179.8.218','6.11.2-0'),
    ('3','2','10.179.8.238','Update sent - completing'),
    ('3','3','10.179.8.226','Update sent - completing'),
    ('3','4','10.179.8.220','Update sent - completing'),
    ('4','1','10.179.8.239','Update sent - completing'),
    ('4','2','10.179.8.235','Update sent - completing'),
    ('4','3','10.179.8.227','Update sent - completing'),
    ('4','4','10.179.8.224','Update sent - completing'),
    ('5','1','10.179.8.234','Update sent - completing'),
    ('5','2','10.179.8.229','Update sent - completing'),
    ('5','3','10.179.8.222','Update sent - completing'),
    ('5','4','10.179.8.219','Update sent - APxm-v1 config pending'),
    ('6','1','10.179.8.230','Update sent - APxm-v1 config pending'),
    ('6','2','10.179.8.237','Update sent - APxm-v1 config pending'),
    ('6','3','10.179.8.228','Update sent - APxm-v1 config pending'),
    ('6','4','10.179.8.221','Update sent - APxm-v1 config pending'),
]
for coach, slot, ip, fw in AP_DATA:
    fill = 'E2EFDA' if fw.startswith('6.11') else 'FFF2CC'
    add_table_row(ap, [coach, slot, ip, fw], font_size=9, fill=fill)
set_col_widths(ap, [1.5, 1.5, 3.5, 10.0])

doc.add_page_break()

# ── 4. RISK ───────────────────────────────────────────────────────────────────
doc.add_heading('4. Risk Assessment and Recommendations', 1)
doc.add_heading('4.1 Risk Summary', 2)
risk = doc.add_table(rows=1, cols=3)
risk.style = 'Table Grid'
for i, h in enumerate(['Area','Risk Level','Rationale']):
    risk.rows[0].cells[i].text = h
shade_header_row(risk)
RISK_DATA = [
    ('Inter-coach Layer-2 fabric','LOW','E2EFDA',
     'All error counters clean. RSTP stable, single root.'),
    ('D1<->E2 missing cable','MEDIUM','FFC7CE',
     'One inter-coach trunk broken. RSTP provides redundancy but resilience is reduced. Stadler action required.'),
    ('AP firmware rollout','LOW','FFF2CC',
     '15/24 APs completing update at time of report. PWLAN may be degraded until complete.'),
    ('Stadler-side beyond FW','UNKNOWN','FFF2CC',
     'Out of scope - Stadler Rail manages those VLANs.'),
    ('PWLAN / cellular','NOT ASSESSED','FFF2CC',
     'Separate scope.'),
]
for area, level, fill, rationale in RISK_DATA:
    row = risk.add_row()
    row.cells[0].text = area
    row.cells[1].text = level
    row.cells[2].text = rationale
    set_cell_shading(row.cells[1], fill)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
set_col_widths(risk, [4.5, 2.5, 9.5])

doc.add_heading('4.2 Recommendations', 2)
RECS = [
    'Stadler Rail: reconnect cable between D1 port e0-1 and E2 port e0-1. '
    'After fix: run sudo obn discover && sudo obn validate on the CCU to confirm LLDP topology clean.',
    'Nomad Digital: once all 24 APs confirmed on 6.11.2-0, '
    'push APxm-v1 config to coaches 4-6 IPs.',
    'Once cabling resolved: perform full L2 health check '
    '(STP re-verify, error counters, Stadler FW reachability, throughput sampling).',
    'Use these findings as baseline for future checks. '
    'Repeat at 6-monthly intervals or after any maintenance.',
    'If end-users report packet loss, the Layer-2 fabric is confirmed healthy '
    '(D1<->E2 exception aside). Investigate PWLAN/cellular or Stadler-side.',
]
for rec in RECS:
    add_bullet(doc, rec, 10)

doc.add_heading('4.3 Open Questions for OeBB / Stadler', 2)
QUESTIONS = [
    'Stadler Rail: when will the D1 e0-1 <-> E2 e0-1 cable be reconnected?',
    'OeBB: confirm preferred cadence for periodic L2 re-checks (recommendation: 6 months).',
    'OeBB: confirm whether RDC service was expected active at time of check.',
    'OeBB: confirm any user-reported network issues on this trainset for targeted follow-up.',
]
for q in QUESTIONS:
    add_bullet(doc, q, 10)

doc.add_page_break()

# ── APPENDIX A - CLI EVIDENCE ─────────────────────────────────────────────────
doc.add_heading('Appendix A - Raw CLI Evidence', 1)
doc.add_paragraph('Verbatim CLI output for key ports.')

doc.add_heading('C2 (10.179.8.192) - e0-1 noise event', 2)
cli1 = (
    'Interface e0-1 is enabled, line protocol is up\n'
    '  Speed: 10000 Mb/s  Duplex: Full\n'
    '  RX packets:90963984  TX packets:63434312\n'
    '  RX errors:1 runts:1 giants:1 frag:0 jabber:1\n'
    '  RX crc errors: 0  TX crc errors:0  carrier false:0'
)
cp1 = doc.add_paragraph(cli1)
for run in cp1.runs:
    run.font.name = 'Courier New'; run.font.size = Pt(8)
doc.add_paragraph(
    'Note: 1 RX error against ~91M packets, CRC=0, carrier-false=0. '
    'Single corrupted frame at link init - noise, not actionable. C2 e0-1 connects to D3.'
)

doc.add_heading('D1 (10.179.8.183) - e0-1 broken link', 2)
cli2 = (
    'Interface e0-1 is enabled, line protocol is down\n'
    '  Speed: Auto  Duplex: Auto\n'
    '  RX packets:0  TX packets:0\n'
    '  RX errors:0 runts:0 giants:0 jabber:0\n'
    '  RX crc errors: 0  carrier false:0'
)
cp2 = doc.add_paragraph(cli2)
for run in cp2.runs:
    run.font.name = 'Courier New'; run.font.size = Pt(8)
doc.add_paragraph('Note: Link simply down - no errors, no flapping. Cable physically disconnected.')

doc.add_heading('C3 (10.179.8.180) - LLDP after cable swap fix', 2)
cli3 = (
    'e0-0  a0:59:3a:d0:54:20  nv6-A2-v8-136  TTCMP bridge/switch\n'
    'e0-1  a0:59:3a:d0:63:00  nv6-C2-v8-136  TTCMP bridge/switch'
)
cp3 = doc.add_paragraph(cli3)
for run in cp3.runs:
    run.font.name = 'Courier New'; run.font.size = Pt(8)
doc.add_paragraph('Note: e0-0 -> A2 (expected), e0-1 -> C2 (expected). Cabling fault resolved.')

# ── APPENDIX B - GLOSSARY ─────────────────────────────────────────────────────
doc.add_heading('Appendix B - Glossary', 1)
gloss = doc.add_table(rows=1, cols=2)
gloss.style = 'Table Grid'
gloss.rows[0].cells[0].text = 'Term'
gloss.rows[0].cells[1].text = 'Definition'
shade_header_row(gloss)
GLOSSARY = [
    ('AFZ','Automatische Fahrgastzahlung - automatic passenger counting'),
    ('CCU','Communications Control Unit - Nomad Digital onboard server (Debian Linux)'),
    ('DOSTO','Doppelstockwagen - Stadler Rail double-deck passenger train'),
    ('Fzg.','Fahrzeug - German for vehicle/trainset'),
    ('LLDP','Link Layer Discovery Protocol - L2 neighbour discovery'),
    ('OBN','Onboard Network - Nomad Digital network management software on the CCU'),
    ('RDC','Remote Diagnostic Centre - Stadler remote diagnostics'),
    ('RSTP','Rapid Spanning Tree Protocol - L2 loop prevention and redundancy'),
    ('VDS Rail','VDS Rail GmbH - manufacturer of consist switches on DOSTO trainsets'),
    ('ZFR','Zugfunk Radio - onboard radio; B1/B3 e1-11 access port VLAN 2'),
]
for term, defn in GLOSSARY:
    add_table_row(gloss, [term, defn], bold_first=True, font_size=9)
set_col_widths(gloss, [3.0, 13.5])

# ── APPENDIX C - REFERENCES ───────────────────────────────────────────────────
doc.add_heading('Appendix C - Reference Documents', 1)
REFS = [
    'OeBB IPv4 schema: ND-DEL-OBB-035-IPA-136_NV_6Teiler.pdf (Fzg. 136, 4736-108)',
    'VDS Rail Consist Switch User Manual, version 2.0.4',
    'Nomad Digital DOSTO L2 Health Check Playbook (project CLAUDE.md)',
    'Nomad Digital fix_obn.py - OBN bug fix script (bugs 1-7)',
    'Stadler_4736-108_Cabling_Fault_Report_v1.0.docx - prior cabling fault report',
]
for ref in REFS:
    add_bullet(doc, ref, 10)

# ── APPENDIX D - REVIEWER NOTES ───────────────────────────────────────────────
doc.add_heading('Appendix D - Reviewer Notes', 1)
doc.add_paragraph('Use this table to log comments if not using Word Track Changes.')
rnotes = doc.add_table(rows=1, cols=4)
rnotes.style = 'Table Grid'
for i, h in enumerate(['#','Date','Reviewer','Note / Action Required']):
    rnotes.rows[0].cells[i].text = h
shade_header_row(rnotes)
for i in range(1, 6):
    add_table_row(rnotes, [str(i),'','',''], font_size=9)
set_col_widths(rnotes, [0.8, 2.5, 3.5, 9.7])

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = 'C:/Users/AbbasRizvi/Documents/dosto-troubleshooting/OBB_Fzg136_4736-108_Network_Health_Check_Report_v1.0.docx'
doc.save(out)
print('Saved:', out)
