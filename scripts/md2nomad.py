#!/usr/bin/env python3
"""Render markdown into the Nomad BMS-ENGI-FOR-006 house template.

Clones ND-DEL-OBB-035_DOSTO_Neu_HW_SW_Release_Note_v1.0.docx (which is the
BMS template already populated for this project), keeps its cover page,
document-control table, disclaimer, TOC field and section/header wiring,
replaces the body, and rewrites the cover fields.

Nomad styles used: Main header Nomad / Header 2 Nomad / Header 3 Nomad /
Header 4 Nomad / Body Nomad / Bullets Nomad / Document control table Nomad.
"""
import re
import sys
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
TEMPLATE = 'ND-DEL-OBB-035_DOSTO_Neu_HW_SW_Release_Note_v1.0.docx'

# body content in the template spans these child indices (inclusive/exclusive)
BODY_START, BODY_END = 50, 145

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')


def add_runs(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        else:
            p.add_run(part)


def parse_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_sep(line):
    return bool(re.fullmatch(r'\|[\s:\-\|]+\|', line.strip()))


def style_or(doc, name, fallback):
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback


def set_cover_field(doc, label, value):
    """Cover fields are 'Label\tValue' paragraphs near the top.

    Run layout in this template is: [label][ ][\t][value][ ][value cont...].
    The tab sits in its own run. So: keep every run up to and including the
    tab (that's the label and its formatting), put the new value in the first
    run after the tab, and blank the remaining runs -- the value is often
    split across several, and leaving any behind renders the old text.
    """
    for p in doc.paragraphs[:16]:
        if not p.runs or not p.text.startswith(label):
            continue
        tab_i = next((j for j, r in enumerate(p.runs) if '\t' in r.text), None)
        if tab_i is None:
            continue
        after = p.runs[tab_i + 1:]
        if after:
            after[0].text = value
            for r in after[1:]:
                r.text = ''
        else:
            # e.g. 'Number of pages \t' with no value run yet
            p.runs[tab_i].text = p.runs[tab_i].text + value
        return True
    return False


def build(md_path, out_path, cover, cover_extra=None):
    cover_extra = cover_extra or {}
    doc = Document(TEMPLATE)
    body = doc.element.body
    els = list(body)

    # strip the old body, keep cover/TOC and trailing sectPr
    for el in els[BODY_START:BODY_END]:
        body.remove(el)

    # anchor: insert before whatever now follows the TOC (the final sectPr)
    tail = list(body)[BODY_START:]
    anchor = tail[0] if tail else None

    S_H1 = style_or(doc, 'Main header Nomad', 'Heading 1')
    S_H2 = style_or(doc, 'Header 2 Nomad', 'Heading 2')
    S_H3 = style_or(doc, 'Header 3 Nomad', 'Heading 3')
    S_H4 = style_or(doc, 'Header 4 Nomad', 'Heading 4')
    S_BODY = style_or(doc, 'Body Nomad', 'Normal')
    S_BUL = style_or(doc, 'Bullets Nomad', 'List Bullet')
    S_TBL = style_or(doc, 'Table Grid4', 'Table Grid')

    def emit(el):
        if anchor is not None:
            anchor.addprevious(el)
        else:
            body.append(el)

    def new_para(style=None):
        p = doc.add_paragraph(style=style)
        el = p._p
        body.remove(el)
        emit(el)
        return p

    lines = open(md_path, encoding='utf-8').read().split('\n')
    i = 0
    skip_doc_details = True
    while i < len(lines):
        s = lines[i].strip()

        # skip the markdown front-matter (cover + doc control + disclaimer):
        # the template already carries these. Body starts at '## 1. Introduction'.
        if skip_doc_details:
            if re.match(r'^##\s+1\.', s):
                skip_doc_details = False
            else:
                i += 1
                continue

        if not s:
            i += 1
            continue

        # fenced image
        if s.startswith('```'):
            spec = s[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            if spec.startswith('img:'):
                parts = spec[4:].split(':')
                path = parts[0]
                width = float(parts[1]) if len(parts) > 1 else 6.2
                p = new_para()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(width))
            continue

        # table
        if s.startswith('|') and i + 1 < len(lines) and is_sep(lines[i + 1]):
            hdr = parse_row(s)
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(parse_row(lines[i]))
                i += 1
            ncol = len(hdr)
            t = doc.add_table(rows=1, cols=ncol)
            try:
                t.style = doc.styles[S_TBL]
            except KeyError:
                pass
            body.remove(t._tbl)
            emit(t._tbl)
            blank_hdr = all(h == '' for h in hdr)
            for j, h in enumerate(hdr):
                cell = t.rows[0].cells[j]
                cell.text = ''
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            # repeat header row across pages
            trPr = t.rows[0]._tr.get_or_add_trPr()
            th = OxmlElement('w:tblHeader')
            th.set(qn('w:val'), 'true')
            trPr.append(th)
            for row in rows:
                cells = t.add_row().cells
                for j in range(min(ncol, len(row))):
                    cells[j].text = ''
                    add_runs(cells[j].paragraphs[0], row[j])
            new_para(S_BODY)
            continue

        m = re.match(r'^(#{1,4})\s+(.*)$', s)
        if m:
            lvl, txt = len(m.group(1)), m.group(2)
            txt = re.sub(r'\*\*(.+?)\*\*', r'\1', txt)
            txt = re.sub(r'\*(.+?)\*', r'\1', txt)
            style = {1: S_H1, 2: S_H2, 3: S_H3, 4: S_H4}[lvl]
            p = new_para(style)
            p.add_run(txt)
            i += 1
            continue

        if s == '---':
            i += 1
            continue

        if re.match(r'^[-*]\s+', s):
            p = new_para(S_BUL)
            add_runs(p, re.sub(r'^[-*]\s+', '', s))
            i += 1
            continue

        if re.match(r'^\d+\.\s+', s):
            p = new_para(style_or(doc, 'Lettered bullets Nomad', 'List Number'))
            add_runs(p, re.sub(r'^\d+\.\s+', '', s))
            i += 1
            continue

        if s.startswith('>'):
            p = new_para(S_BODY)
            add_runs(p, s.lstrip('> ').strip())
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        p = new_para(S_BODY)
        add_runs(p, s)
        i += 1

    for label, value in cover.items():
        ok = set_cover_field(doc, label, value)
        if not ok:
            print(f'  ! cover field not matched: {label}')

    # document-control table: row 1 is the v1 record inherited from the
    # template's own release note -- rewrite it for this document.
    dc = next((t for t in doc.tables
               if t.style and t.style.name == 'Document control table Nomad'), None)
    if dc is not None and len(dc.rows) > 1 and 'doc_control' in cover_extra:
        vals = cover_extra['doc_control']
        for j, v in enumerate(vals):
            if j >= len(dc.rows[1].cells):
                break
            cell = dc.rows[1].cells[j]
            cell.text = ''
            cell.paragraphs[0].add_run(v)

    doc.save(out_path)
    print('wrote', out_path)


if __name__ == '__main__':
    import json
    cfg = json.loads(sys.argv[3]) if len(sys.argv) > 3 else {}
    extra = {}
    if 'doc_control' in cfg:
        extra['doc_control'] = cfg.pop('doc_control')
    build(sys.argv[1], sys.argv[2], cfg, extra)
